"""
tools/resume_formatter.py
--------------------------
Generates a professional, ATS-optimized DOCX resume.
Clean single-column layout: no tables, no text boxes, no images.
Professional blue accent theme with proper spacing.
"""

import os
import re
from datetime import datetime
from typing import Dict, List

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

OUTPUTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "outputs")

# Color palette
BLUE_DARK   = RGBColor(0x1E, 0x3A, 0x5F)   # #1e3a5f — headings
BLUE_MID    = RGBColor(0x2E, 0x74, 0xB5)   # #2e74b5 — rule lines
GRAY_DARK   = RGBColor(0x26, 0x26, 0x26)   # #262626 — body text
GRAY_MID    = RGBColor(0x55, 0x55, 0x55)   # #555555 — secondary text
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


def _set_font(run, name="Calibri", size=10.5, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _set_para_spacing(para, before=0, after=4, line_spacing=None):
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after  = Pt(after)
    if line_spacing:
        from docx.shared import Pt as DPt
        fmt.line_spacing = DPt(line_spacing)


def _add_rule(doc, color_hex="2E74B5"):
    """Add a thin colored horizontal rule under a paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_section_heading(doc, title: str):
    """Add a bold blue section heading with a rule underneath."""
    p = doc.add_paragraph()
    _set_para_spacing(p, before=10, after=0)
    run = p.add_run(title.upper())
    _set_font(run, size=11, bold=True, color=BLUE_DARK)
    _add_rule(doc)


def generate_ats_resume(resume_data: Dict, output_filename: str = None) -> str:
    """
    Generate a professional ATS-optimized DOCX resume.
    Returns path to the generated file.
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx required. Run: pip install python-docx")

    os.makedirs(OUTPUTS_DIR, exist_ok=True)
    if not output_filename:
        output_filename = f"ATS_Resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    output_path = os.path.join(OUTPUTS_DIR, f"{output_filename}.docx")

    doc = Document()

    # ── Page setup ──
    for sec in doc.sections:
        sec.top_margin    = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin   = Inches(0.85)
        sec.right_margin  = Inches(0.85)

    # ── Remove default paragraph spacing ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = GRAY_DARK

    # ════════════════════════════════════════
    # HEADER — Name
    # ════════════════════════════════════════
    name = resume_data.get("name", "Your Name").strip() or "Your Name"
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(name_p, before=0, after=4)
    name_run = name_p.add_run(name.upper())
    _set_font(name_run, size=20, bold=True, color=BLUE_DARK)

    # ── Contact line ──
    contact = resume_data.get("contact", {})
    parts = []
    if contact.get("email"):    parts.append(f"✉ {contact['email']}")
    if contact.get("phone"):    parts.append(f"📱 {contact['phone']}")
    if contact.get("location"): parts.append(f"📍 {contact['location']}")
    if contact.get("linkedin"): parts.append(f"🔗 {contact['linkedin']}")
    if contact.get("github"):   parts.append(f"💻 {contact['github']}")

    if parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(contact_p, before=0, after=2)
        contact_run = contact_p.add_run("  |  ".join(parts))
        _set_font(contact_run, size=9.5, color=GRAY_MID)

    # Full-width blue rule under header
    _add_rule(doc, "1E3A5F")

    # ════════════════════════════════════════
    # PROFESSIONAL SUMMARY
    # ════════════════════════════════════════
    summary = resume_data.get("summary", "").strip()
    if summary:
        _add_section_heading(doc, "Professional Summary")
        p = doc.add_paragraph()
        _set_para_spacing(p, before=0, after=8)
        r = p.add_run(summary)
        _set_font(r, size=10.5, italic=False, color=GRAY_DARK)

    # ════════════════════════════════════════
    # TECHNICAL SKILLS
    # ════════════════════════════════════════
    skills = resume_data.get("skills", [])
    if skills:
        _add_section_heading(doc, "Technical Skills")
        skills_list = skills if isinstance(skills, list) else [s.strip() for s in str(skills).split(",")]
        skills_text = "  •  ".join(s.strip() for s in skills_list if s.strip())
        p = doc.add_paragraph()
        _set_para_spacing(p, before=0, after=8)
        r = p.add_run(skills_text)
        _set_font(r, size=10.5, color=GRAY_DARK)

    # ════════════════════════════════════════
    # PROFESSIONAL EXPERIENCE
    # ════════════════════════════════════════
    experience = resume_data.get("experience", [])
    if experience:
        _add_section_heading(doc, "Professional Experience")
        for job in experience:
            # Company | Date on same line
            header_p = doc.add_paragraph()
            _set_para_spacing(header_p, before=4, after=0)
            company_run = header_p.add_run(job.get("company", ""))
            _set_font(company_run, size=11, bold=True, color=BLUE_DARK)
            if job.get("dates"):
                date_run = header_p.add_run(f"  —  {job['dates']}")
                _set_font(date_run, size=10, color=GRAY_MID)

            # Role title
            role_p = doc.add_paragraph()
            _set_para_spacing(role_p, before=0, after=2)
            role_run = role_p.add_run(job.get("title", ""))
            _set_font(role_run, size=10.5, bold=True, italic=True, color=GRAY_DARK)

            # Bullet points
            for bullet in job.get("bullets", []):
                b = bullet.strip().lstrip("•-– ")
                if not b:
                    continue
                bp = doc.add_paragraph(style="List Bullet")
                _set_para_spacing(bp, before=0, after=2)
                br = bp.add_run(b)
                _set_font(br, size=10.5, color=GRAY_DARK)

        doc.add_paragraph()  # spacer

    # ════════════════════════════════════════
    # EDUCATION
    # ════════════════════════════════════════
    education = resume_data.get("education", [])
    if education:
        _add_section_heading(doc, "Education")
        for edu in education:
            ep = doc.add_paragraph()
            _set_para_spacing(ep, before=4, after=2)
            deg_run = ep.add_run(f"{edu.get('degree','')} — {edu.get('institution','')}")
            _set_font(deg_run, size=11, bold=True, color=BLUE_DARK)
            if edu.get("year"):
                yr_run = ep.add_run(f"  |  {edu['year']}")
                _set_font(yr_run, size=10, color=GRAY_MID)
            if edu.get("details"):
                dp = doc.add_paragraph()
                _set_para_spacing(dp, before=0, after=2)
                dr = dp.add_run(edu["details"])
                _set_font(dr, size=10.5, color=GRAY_DARK)

    # ════════════════════════════════════════
    # KEY PROJECTS
    # ════════════════════════════════════════
    projects = resume_data.get("projects", [])
    if projects:
        _add_section_heading(doc, "Key Projects")
        for proj in projects:
            pp = doc.add_paragraph()
            _set_para_spacing(pp, before=4, after=2)
            name_run = pp.add_run(f"{proj.get('name','')}: ")
            _set_font(name_run, size=11, bold=True, color=BLUE_DARK)
            desc_run = pp.add_run(proj.get("description",""))
            _set_font(desc_run, size=10.5, color=GRAY_DARK)
            if proj.get("tech"):
                tp = doc.add_paragraph()
                _set_para_spacing(tp, before=0, after=4)
                tr = tp.add_run(f"Tech Stack: {proj['tech']}")
                _set_font(tr, size=10, italic=True, color=GRAY_MID)

    # ════════════════════════════════════════
    # CERTIFICATIONS
    # ════════════════════════════════════════
    certs = resume_data.get("certifications", [])
    if certs:
        _add_section_heading(doc, "Certifications")
        for cert in certs:
            cp = doc.add_paragraph(style="List Bullet")
            _set_para_spacing(cp, before=0, after=2)
            cr = cp.add_run(cert)
            _set_font(cr, size=10.5, color=GRAY_DARK)

    doc.save(output_path)
    return output_path


def build_resume_data_from_llm_output(llm_output: str, user_name: str = "", contact: Dict = None) -> Dict:
    """
    Parse LLM-generated resume text into structured resume_data dict.
    """
    resume_data = {
        "name": user_name or "Your Name",
        "contact": contact or {},
        "summary": "",
        "skills": [],
        "experience": [],
        "education": [],
        "projects": [],
        "certifications": [],
        "llm_output": llm_output
    }

    # ── Extract SUMMARY ──
    m = re.search(r"(?:PROFESSIONAL SUMMARY|SUMMARY|OBJECTIVE|PROFILE)[:\n\s]+(.*?)(?:\n[A-Z]{3,}[\s\n]|\Z)",
                  llm_output, re.DOTALL | re.IGNORECASE)
    if m:
        resume_data["summary"] = m.group(1).strip()[:600]
    else:
        paras = [p.strip() for p in llm_output.split("\n\n") if p.strip() and len(p.strip()) > 50]
        if paras:
            resume_data["summary"] = paras[0][:400]

    # ── Extract SKILLS ──
    m = re.search(r"(?:TECHNICAL SKILLS|SKILLS|CORE COMPETENCIES)[:\n\s]+(.*?)(?:\n[A-Z]{3,}[\s\n]|\Z)",
                  llm_output, re.DOTALL | re.IGNORECASE)
    if m:
        raw = m.group(1).strip()
        skills = [s.strip().lstrip("•-– ") for s in re.split(r"[,•|\n]+", raw) if s.strip() and len(s.strip()) > 1]
        resume_data["skills"] = skills[:20]

    # ── Extract EXPERIENCE ──
    exp_section = re.search(
        r"(?:PROFESSIONAL EXPERIENCE|EXPERIENCE|WORK EXPERIENCE)[:\n\s]+(.*?)(?:\nEDUCATION|\nKEY PROJECTS|\nPROJECTS|\nCERTIFICATION|\Z)",
        llm_output, re.DOTALL | re.IGNORECASE
    )
    if exp_section:
        exp_text = exp_section.group(1).strip()
        # Split by company blocks (lines that look like "Company | Role | Dates" or "Company  —  Date")
        blocks = re.split(r"\n(?=[A-Z][^\n]{0,60}(?:\||\—|—)[^\n]{0,60}\n)", exp_text)
        for block in blocks:
            lines = [l.strip() for l in block.strip().split("\n") if l.strip()]
            if not lines:
                continue
            bullets = [l.lstrip("•●*-– ") for l in lines[2:] if l.startswith(("•", "●", "-", "–", "*", " "))]
            job = {
                "company": lines[0] if lines else "",
                "title":   lines[1] if len(lines) > 1 else "",
                "dates":   "",
                "bullets": bullets if bullets else lines[2:6]
            }
            if job["company"] or job["title"]:
                resume_data["experience"].append(job)

    return resume_data
